#!/usr/bin/env python
# encoding: utf-8
# Date: 2023/02/15
# file: _docx.py
# Email:
# Author: rakiwine

"""
八号  5
七号  5.5
小六  6.5
六号  7.5
小五  9
五号  10.5
小四  12
四号  14
小三  16
三号  18
小二  20
二号  22
小一  24
一号  26
小初  36
初号  42
"""

import re
import docx
from docx.enum.style import WD_STYLE_TYPE
# WD_ORIENT 改 WD_ORIENTATION
from docx.enum.section import WD_ORIENTATION
# WD_ALIGN_PARAGRAPH 改 WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

# 打开已有docx
existDocument = docx.Document("绝对路径")

# WD_STYLE_TYPE 段落样式（PARAGRAPH）、字符样式（CHARACTER）、表格样式（TABLE）和列表样式（LIST）
# 给当前文档增加设定样式
redStyle = existDocument.styles.add_style('redStyle', WD_STYLE_TYPE.PARAGRAPH)
redStyle.font.name = '宋体(正文)'
redStyle.font.size = Pt(12)
redStyle.font.color.rgb = RGBColor(255, 0, 0)
redStyle.font.color.rgb = RGBColor.from_string("FF0000")
"""
微软雅黑
宋体
宋体(正文)
"""
redStyle.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体(正文)')

# 获取所有段落列表
paragraphList = existDocument.paragraphs

for paragraph in paragraphList:
    # 获取锻炼文字
    lineDeal = paragraph.text

    # 段落包含换行符 再分割为行
    for line in lineDeal.split("\n"):
        pass

# 创建一个空白docx
newDocument = docx.Document()
# 设置全文西文字体
newDocument.styles['Normal'].font.name = 'Calibri'
# 设置全文西文字体大小
newDocument.styles['Normal'].font.size = Pt(12)
# 设置全文中文字体
newDocument.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体(正文)')


def _layout_():
    # 获取文档章节列表，新建的文档只有一节
    sections = newDocument.sections
    section_0 = sections[0]

    # 获取第一个节的页眉
    header = section_0.header
    footer = section_0.footer
    header.add_paragraph('这是第一节的页眉')
    footer .add_paragraph('这是第一节的页脚')

    # 新建文档默认为竖向，改为横向须设置三个参数：
    # 若只设置第一个而不设置另外两个，则页面方向不发生变化；
    # 如果只设置后两个，则页面的宽和高会改变，但仍为纵向
    # 设置页面方向为横向须同时设置高度和宽度才有用
    section_0.orientation = WD_ORIENTATION.LANDSCAPE
    section_0.page_height = Cm(21)
    section_0.page_width = Cm(29.7)

    # 设置页边距均为1cm
    section_0.left_margin = Cm(1)
    section_0.right_margin = Cm(1)
    section_0.top_margin = Cm(1)
    section_0.bottom_margin = Cm(1)


def _heading_():
    # level 范围为 0 至 9 默认1
    _heading = newDocument.add_heading(u"", level=1)


def _paragraph_():
    # 添加段落 设定预留样式
    _paragraph = newDocument.add_paragraph("", style=None)
    _paragraph.insert_paragraph_before(u"", style=None)

    _paragraph.paragraph_format.line_spacing = Pt(25)  # 设置行高

    # 设置文字在这一行的纵向位置，通常设置段后为0，和行号，来调整行间距
    # a.paragraph_format.space_before = Pt(0)               # 段前
    _paragraph.paragraph_format.space_after = Pt(0)         # 段后

    # a.paragraph_format.left_indent = Pt(20)               #左缩进
    # a.paragraph_format.right_indent = Pt(20)              #右缩进
    _paragraph.paragraph_format.first_line_indent = Pt(20)  # 首行缩进

    # 添加段落 并设定段落样式
    _paragraph = newDocument.add_paragraph()
    _paragraph = newDocument.paragraphs[-1]
    _text = _paragraph.add_run("第一段文字是中文；The first paragraph is in Chinese")
    _text.font.size = Pt(15)
    _text.bold = True
    _text.font.name = '宋体(正文)'
    _text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体(正文)')


def _picture_():
    # 添加图片 只设定一个方向的尺寸，另一方向会自动缩放
    newDocument.add_picture('img.png', width=Cm(10))


def _table_():
    """
    _TableStyle('Normal Table') id: 187730312
    _TableStyle('Table Grid') id: 187730312
    _TableStyle('Light Shading') id: 187730376
    _TableStyle('Light Shading Accent 1') id: 187730312
    _TableStyle('Light Shading Accent 2') id: 187730376
    _TableStyle('Light Shading Accent 3') id: 187730312
    _TableStyle('Light Shading Accent 4') id: 187730376
    _TableStyle('Light Shading Accent 5') id: 187730312
    _TableStyle('Light Shading Accent 6') id: 187730376
    _TableStyle('Light List') id: 187730312
    :return:
    """
    # 添加表格 五行五列
    table = newDocument.add_table(5, 5, style='Table Grid')
    # len(table.rows), len(table.columns)

    # 在第一行第一列的单元格填充字符1
    # 表格的行列都从0开始计数
    table.cell(0, 0).text = '1'

    # 行列的索引
    row = table.rows[1]
    column = table.columns[1]
    row0_cells = table.row_cells(0)
    col0_cells = table.column_cells(0)

    # 按行遍历单元格
    for row in table.rows:
        for cell in row.cells:
            # print(cell)
            pass
    # 按列遍历单元格
    for col in table.columns:
        for cell in col.cells:
            # print(cell)
            pass

    # 合并单元格
    table.cell(2, 1).merge(table.cell(3, 2))

    # 添加行
    table.add_row()
    # 添加列需要指定宽度
    table.add_column(Cm(5))
    # 行的删除：可通过row._element.getparent().remove()函数实现
    row._element.getparent().remove(row._element)

    # 列的删除：列没有像行那样的remove()函数，可通过遍历列中的单元格进行删除.
    # 但表格中的cell按行存储，删除第i行的某个单元格后，每行的cell数并不变化，
    # 逻辑上第i+1行的第一个单元格会补到第i行的最后一个单元格，所以与外观不同。
    for cell in column.cells:
        cell._element.getparent().remove(cell._element)

    # 设置行高列宽
    table.rows[1].height = Cm(1)
    table.columns[1].height = Cm(1)
    # 同列单元格宽度相同，不同时以最大宽度为准
    table.cell(1, 1).width = Cm(3)

    # 单元格中第一段落 居中
    # WD_ALIGN_PARAGRAPH 改 WD_PARAGRAPH_ALIGNMENT
    # 水平居中|左对齐|右对齐|两端对齐
    # WD_ALIGN_PARAGRAPH.CENTER | WD_ALIGN_PARAGRAPH.LEFT | WD_ALIGN_PARAGRAPH.RIGHT | WD_ALIGN_PARAGRAPH.DISTRIBUTE
    # 垂直居中|顶部对齐|底部对齐
    # WD_ALIGN_VERTICAL.CENTER | WD_ALIGN_VERTICAL.TOP | WD_ALIGN_VERTICAL.BOTTOM
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.cell(0, 0).paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.TOP

    # 覆盖
    table.cell(0, 2).text = ""
    table.cell(1, 2).paragraphs[0].add_run("")

    _paragraph_cell = table.cell(0, 1).paragraphs[-1].add_run("")
    _paragraph_cell.font.name = '华文琥珀'
    _paragraph_cell.font.size = Pt(30)
    _paragraph_cell.element.rPr.rFonts.set(qn('w:eastAsia'), '华文行楷')

    table.cell(0, 1).paragraphs[-1].add_run().add_picture(r"img.png", width=Cm(5))

    _paragraph = table.cell(0, 1).add_paragraph()
    _paragraph = table.cell(0, 1).paragraphs[-1]
    _text = _paragraph.add_run()
    _text.add_picture(r"img.png", width=Cm(5))


# 段落列表最后一段 居中
newDocument.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
newDocument.paragraphs[-1].vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


newDocument.save("绝对路径")
